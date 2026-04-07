"""
MD → DOCX 변환 스크립트
마크다운 파일을 python-docx를 사용하여 Word 문서로 변환합니다.
정규식 기반 상태 머신으로 코드 블록, 테이블, 리스트를 처리합니다.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from style_config import (
  FONT_BODY, FONT_CODE, FONT_CODE_FALLBACK,
  FONT_SIZE_BODY, FONT_SIZE_H1, FONT_SIZE_H2, FONT_SIZE_H3, FONT_SIZE_H4,
  FONT_SIZE_CODE, FONT_SIZE_FOOTER,
  COLOR_PRIMARY, COLOR_SECONDARY, COLOR_ACCENT, COLOR_CODE_BG,
  COLOR_TABLE_HEADER, COLOR_TABLE_HEADER_TEXT, COLOR_TABLE_ROW_ALT,
  MARGIN_TOP_CM, MARGIN_BOTTOM_CM, MARGIN_LEFT_CM, MARGIN_RIGHT_CM,
  LINE_SPACING,
)


def hexToRgb(hexColor: str) -> RGBColor:
  """HEX 색상을 RGBColor로 변환"""
  hexColor = hexColor.lstrip("#")
  return RGBColor(
    int(hexColor[0:2], 16),
    int(hexColor[2:4], 16),
    int(hexColor[4:6], 16),
  )


def setupDocument(doc: Document):
  """문서 기본 설정 (여백, 폰트)"""
  section = doc.sections[0]
  section.top_margin = Cm(MARGIN_TOP_CM)
  section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
  section.left_margin = Cm(MARGIN_LEFT_CM)
  section.right_margin = Cm(MARGIN_RIGHT_CM)

  # 기본 폰트 설정
  style = doc.styles["Normal"]
  font = style.font
  font.name = FONT_BODY
  font.size = Pt(FONT_SIZE_BODY)
  style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

  # 줄 간격
  paragraphFormat = style.paragraph_format
  paragraphFormat.line_spacing = LINE_SPACING


def addPageNumbers(doc: Document):
  """페이지 번호를 푸터에 추가"""
  section = doc.sections[0]
  footer = section.footer
  footer.is_linked_to_previous = False

  paragraph = footer.paragraphs[0]
  paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

  run = paragraph.add_run()
  run.font.size = Pt(FONT_SIZE_FOOTER)
  run.font.name = FONT_BODY

  # 페이지 번호 필드 삽입
  fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
  run._element.append(fldChar1)

  instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
  run._element.append(instrText)

  fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
  run._element.append(fldChar2)


def addHeading(doc: Document, text: str, level: int):
  """헤딩 추가 (한국어 폰트 적용)"""
  heading = doc.add_heading(text, level=level)

  sizeMap = {1: FONT_SIZE_H1, 2: FONT_SIZE_H2, 3: FONT_SIZE_H3, 4: FONT_SIZE_H4}
  colorMap = {1: COLOR_PRIMARY, 2: COLOR_SECONDARY, 3: COLOR_ACCENT, 4: COLOR_ACCENT}

  for run in heading.runs:
    run.font.name = FONT_BODY
    run.font.size = Pt(sizeMap.get(level, FONT_SIZE_H4))
    run.font.color.rgb = hexToRgb(colorMap.get(level, COLOR_ACCENT))
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)


def addCodeBlock(doc: Document, codeLines: list[str], language: str = ""):
  """코드 블록 추가 (회색 배경 + 고정폭 폰트)"""
  codeText = "\n".join(codeLines)

  paragraph = doc.add_paragraph()
  paragraph.paragraph_format.space_before = Pt(6)
  paragraph.paragraph_format.space_after = Pt(6)

  # 배경색 설정
  shading = parse_xml(
    f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>'
  )
  paragraph.paragraph_format.element.get_or_add_pPr().append(shading)

  run = paragraph.add_run(codeText)
  run.font.name = FONT_CODE
  run.font.size = Pt(FONT_SIZE_CODE)
  run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CODE_FALLBACK)


def addTable(doc: Document, rows: list[list[str]]):
  """마크다운 테이블을 Word 테이블로 변환"""
  if not rows or len(rows) < 1:
    return

  numCols = len(rows[0])
  table = doc.add_table(rows=len(rows), cols=numCols)
  table.alignment = WD_TABLE_ALIGNMENT.CENTER
  table.style = "Table Grid"

  for i, row in enumerate(rows):
    for j, cellText in enumerate(row):
      if j >= numCols:
        break
      cell = table.cell(i, j)
      cell.text = cellText.strip()

      # 첫 행 헤더 스타일
      if i == 0:
        shading = parse_xml(
          f'<w:shd {nsdecls("w")} w:fill="{COLOR_TABLE_HEADER.lstrip("#")}" w:val="clear"/>'
        )
        cell._element.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
          for run in paragraph.runs:
            run.font.color.rgb = hexToRgb(COLOR_TABLE_HEADER_TEXT)
            run.font.bold = True
            run.font.name = FONT_BODY
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
      # 짝수행 배경
      elif i % 2 == 0:
        shading = parse_xml(
          f'<w:shd {nsdecls("w")} w:fill="{COLOR_TABLE_ROW_ALT.lstrip("#")}" w:val="clear"/>'
        )
        cell._element.get_or_add_tcPr().append(shading)


def addFormattedParagraph(doc: Document, text: str):
  """인라인 포맷 (볼드, 이탤릭, 코드, 링크) 처리"""
  paragraph = doc.add_paragraph()

  # 인라인 패턴 분리
  # **bold**, *italic*, `code`, [text](url)
  pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[(.+?)\]\((.+?)\))"
  lastEnd = 0

  for match in re.finditer(pattern, text):
    # 매치 전 일반 텍스트
    if match.start() > lastEnd:
      run = paragraph.add_run(text[lastEnd:match.start()])
      run.font.name = FONT_BODY
      run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

    if match.group(2):  # **bold**
      run = paragraph.add_run(match.group(2))
      run.bold = True
      run.font.name = FONT_BODY
      run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    elif match.group(3):  # *italic*
      run = paragraph.add_run(match.group(3))
      run.italic = True
      run.font.name = FONT_BODY
      run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    elif match.group(4):  # `code`
      run = paragraph.add_run(match.group(4))
      run.font.name = FONT_CODE
      run.font.size = Pt(FONT_SIZE_CODE)
    elif match.group(5):  # [text](url)
      run = paragraph.add_run(f"{match.group(5)} ({match.group(6)})")
      run.font.name = FONT_BODY
      run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

    lastEnd = match.end()

  # 나머지 텍스트
  if lastEnd < len(text):
    run = paragraph.add_run(text[lastEnd:])
    run.font.name = FONT_BODY
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)


def preprocessForDocx(mdContent: str) -> str:
  """DOCX 변환 전 포맷별 후처리.

  - 앵커 링크 `[text](#anchor)`는 DOCX에서 페이지 점프가 동작하지 않으므로
    `text`만 남김 (목차의 잡다한 `#anchor` 노출 방지)
  - HTML 주석 `<!-- ... -->` 제거
  - HTML 인라인 태그 제거 (가독성 깨짐 방지)
  """
  # 앵커 링크: [text](#anchor) → text
  mdContent = re.sub(r"\[([^\]]+)\]\(#[^)]+\)", r"\1", mdContent)
  # HTML 주석 제거
  mdContent = re.sub(r"<!--.*?-->", "", mdContent, flags=re.DOTALL)
  # 단독 HTML 태그(예: <br>, <div>) 제거 — 코드 블록 밖만 영향
  mdContent = re.sub(r"<(br|hr|div|span)\s*/?>", "", mdContent, flags=re.IGNORECASE)
  return mdContent


def convertMdToDocx(mdContent: str, outputPath: Path):
  """마크다운을 DOCX로 변환하는 메인 함수"""
  mdContent = preprocessForDocx(mdContent)

  doc = Document()
  setupDocument(doc)
  addPageNumbers(doc)

  lines = mdContent.split("\n")
  i = 0
  inCodeBlock = False
  codeLines = []
  codeLang = ""
  inTable = False
  tableRows = []

  while i < len(lines):
    line = lines[i]

    # --- 코드 블록 처리 ---
    if line.strip().startswith("```"):
      if inCodeBlock:
        # 코드 블록 종료
        addCodeBlock(doc, codeLines, codeLang)
        codeLines = []
        codeLang = ""
        inCodeBlock = False
      else:
        # 코드 블록 시작
        inCodeBlock = True
        codeLang = line.strip().lstrip("`").strip()
      i += 1
      continue

    if inCodeBlock:
      codeLines.append(line)
      i += 1
      continue

    # --- 테이블 처리 ---
    if "|" in line and line.strip().startswith("|"):
      stripped = line.strip()
      # 구분선 행인지 확인 (|---|---|)
      if re.match(r"^\|[\s\-:|]+\|$", stripped):
        i += 1
        continue

      cells = [c.strip() for c in stripped.split("|")[1:-1]]
      if not inTable:
        inTable = True
        tableRows = []
      tableRows.append(cells)
      i += 1
      continue
    elif inTable:
      # 테이블 종료
      addTable(doc, tableRows)
      tableRows = []
      inTable = False

    # --- 헤딩 ---
    headingMatch = re.match(r"^(#{1,4})\s+(.+)$", line)
    if headingMatch:
      level = len(headingMatch.group(1))
      text = headingMatch.group(2).strip()
      addHeading(doc, text, level)
      i += 1
      continue

    # --- 수평선 ---
    if re.match(r"^---+$", line.strip()):
      doc.add_paragraph("─" * 50)
      i += 1
      continue

    # --- 인용 ---
    if line.strip().startswith(">"):
      quoteText = line.strip().lstrip("> ").strip()
      paragraph = doc.add_paragraph()
      paragraph.paragraph_format.left_indent = Cm(1.0)
      run = paragraph.add_run(f"│ {quoteText}")
      run.italic = True
      run.font.name = FONT_BODY
      run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
      i += 1
      continue

    # --- 리스트 (순서 없음) ---
    listMatch = re.match(r"^(\s*)[-*]\s+(.+)$", line)
    if listMatch:
      indent = len(listMatch.group(1)) // 2
      text = listMatch.group(2)
      bullet = "  " * indent + "• " + text
      addFormattedParagraph(doc, bullet)
      i += 1
      continue

    # --- 리스트 (순서 있음) ---
    orderedMatch = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
    if orderedMatch:
      text = orderedMatch.group(2)
      addFormattedParagraph(doc, text)
      i += 1
      continue

    # --- 빈 줄 ---
    if not line.strip():
      i += 1
      continue

    # --- 일반 텍스트 ---
    addFormattedParagraph(doc, line)
    i += 1

  # 남은 테이블 처리
  if inTable and tableRows:
    addTable(doc, tableRows)

  # 저장
  outputPath.parent.mkdir(parents=True, exist_ok=True)
  doc.save(str(outputPath))
  print(f"  [완료] DOCX 생성: {outputPath}")


def main():
  """메인 실행 함수"""
  scriptDir = Path(__file__).parent
  projectRoot = scriptDir.parent
  mdPath = projectRoot / "output" / "Guidebook.md"
  docxPath = projectRoot / "output" / "Guidebook.docx"

  print("=" * 60)
  print("  MD → DOCX 변환 시작")
  print("=" * 60)

  if not mdPath.exists():
    print(f"[오류] 통합 MD 파일이 없습니다: {mdPath}")
    print("  먼저 md_merger.py를 실행해주세요.")
    return

  mdContent = mdPath.read_text(encoding="utf-8")
  convertMdToDocx(mdContent, docxPath)


if __name__ == "__main__":
  main()
